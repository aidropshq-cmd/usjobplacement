"""Resume parsing.

Runs entirely in this process. `pypdf` and `python-docx` extract the text and
plain rules extract the fields — no model is called and no bytes leave our
infrastructure, which is what the site's privacy note promises people.

The important design decision is that **every field carries a confidence and
nothing is written to a profile automatically**. Extraction lands in
ResumeExtraction as a suggestion; the candidate confirms or corrects it. A
regex's opinion of somebody's career must never silently replace their own.

Confidence values are set by how reliable the rule genuinely is, not by
wishful thinking:

    0.95  email, phone       — unambiguous patterns
    0.80  skills             — matched against a curated vocabulary
    0.70  education, certs   — distinctive keywords
    0.55  years experience   — inferred from date ranges, easily fooled
    0.40  name               — heuristic; resumes have no fixed layout
    0.35  job titles         — matched against common patterns
    0.30  companies          — genuinely hard without a model. Suggestion only.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import date

logger = logging.getLogger(__name__)


class ParseError(RuntimeError):
    """The file could not be read as text."""


# --------------------------------------------------------------- text extraction


def extract_text(data: bytes, content_type: str) -> str:
    if "pdf" in content_type:
        return _pdf_text(data)
    if "wordprocessingml" in content_type or "docx" in content_type:
        return _docx_text(data)
    raise ParseError(f"Cannot read {content_type}")


def _pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        # A scanned resume is an image with no text layer. That is a real and
        # common case, and it must surface as "we could not read this" rather
        # than as an empty profile.
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001 - any pypdf failure is a parse failure
        raise ParseError(f"PDF could not be read: {exc}") from exc


def _docx_text(data: bytes) -> str:
    try:
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        # Plenty of resumes lay everything out in a table.
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise ParseError(f"DOCX could not be read: {exc}") from exc


# ------------------------------------------------------------------- vocabulary

SKILLS = {
    "languages": [
        "Python", "Java", "JavaScript", "TypeScript", "Go", "Golang", "Rust",
        "C++", "C#", ".NET", "Ruby", "PHP", "Scala", "Kotlin", "Swift", "R",
        "SQL", "Bash", "PowerShell", "Apex",
    ],
    "cloud": [
        "AWS", "Azure", "GCP", "Google Cloud", "EC2", "S3", "Lambda", "EKS",
        "ECS", "RDS", "CloudFront", "Route 53", "IAM", "CloudFormation",
        "Azure DevOps", "BigQuery", "Redshift",
    ],
    "devops": [
        "Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins", "GitLab CI",
        "GitHub Actions", "CircleCI", "ArgoCD", "Helm", "Prometheus", "Grafana",
        "Datadog", "Splunk", "CI/CD", "Linux", "Nginx",
    ],
    "data": [
        "Spark", "Hadoop", "Kafka", "Airflow", "dbt", "Snowflake", "Databricks",
        "Pandas", "NumPy", "TensorFlow", "PyTorch", "scikit-learn", "Tableau",
        "Power BI", "ETL", "Machine Learning",
    ],
    "web": [
        "React", "Angular", "Vue", "Next.js", "Node.js", "Django", "Flask",
        "FastAPI", "Spring", "Spring Boot", "Express", "GraphQL", "REST API",
        "HTML", "CSS", "Tailwind",
    ],
    "databases": [
        "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch", "Oracle",
        "DynamoDB", "Cassandra", "SQL Server",
    ],
    "practices": [
        "Agile", "Scrum", "Kanban", "TDD", "Microservices", "Serverless",
        "Salesforce", "Selenium", "Cypress", "JUnit", "pytest", "Jira",
    ],
}

CERTIFICATIONS = [
    "AWS Certified Solutions Architect", "AWS Certified Developer",
    "AWS Certified DevOps Engineer", "AWS Certified SysOps",
    "Azure Solutions Architect", "Azure Administrator", "Azure Developer",
    "Google Cloud Professional", "Certified Kubernetes Administrator", "CKA",
    "CKAD", "Terraform Associate", "PMP", "CISSP", "CompTIA Security+",
    "Salesforce Certified", "Scrum Master", "CSM", "ITIL",
]

DEGREES = [
    "Bachelor", "Master", "PhD", "Ph.D", "Doctorate", "B.S.", "BS", "B.Sc",
    "M.S.", "MS", "M.Sc", "MBA", "B.Tech", "M.Tech", "B.E.", "M.E.",
    "Associate Degree",
]

TITLE_WORDS = [
    "Software Engineer", "Senior Software Engineer", "Staff Engineer",
    "DevOps Engineer", "Site Reliability Engineer", "SRE", "Cloud Engineer",
    "Cloud Architect", "Solutions Architect", "Data Engineer", "Data Scientist",
    "Data Analyst", "Machine Learning Engineer", "QA Engineer",
    "Test Engineer", "Automation Engineer", "Security Engineer",
    "Cybersecurity Analyst", "Salesforce Developer", "Salesforce Administrator",
    "Business Analyst", "Product Manager", "Project Manager",
    "Full Stack Developer", "Backend Developer", "Frontend Developer",
    "Platform Engineer", "Systems Engineer", "Database Administrator",
]

EMAIL_RE = re.compile(r"\b[\w.+-]+@[\w-]+\.[\w.-]{2,}\b")
PHONE_RE = re.compile(r"(?:\+?1[\s.-]?)?\(?([2-9]\d{2})\)?[\s.-]?([2-9]\d{2})[\s.-]?(\d{4})\b")
YEAR_RANGE_RE = re.compile(
    r"(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?"
    r"(19[89]\d|20[0-4]\d)\s*(?:-|–|—|to)\s*"
    r"(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+)?"
    r"(19[89]\d|20[0-4]\d|[Pp]resent|[Cc]urrent|[Nn]ow)",
)
STATED_YEARS_RE = re.compile(r"\b(\d{1,2})\+?\s*(?:years?|yrs?)\b", re.I)

# Lines that are section headers or contact rows, never a person's name.
NOT_A_NAME = re.compile(
    r"@|\d|resume|curriculum|vitae|summary|objective|profile|experience|"
    r"education|skills|contact|linkedin|github|http",
    re.I,
)


def _all_skills() -> list[str]:
    return [skill for group in SKILLS.values() for skill in group]


# ------------------------------------------------------------------ extraction


def _find_skills(text: str) -> list[str]:
    found = []
    for skill in _all_skills():
        # Word boundaries so "R" does not match every capital R, and "Go"
        # does not match "Google". Escaped because of C++, .NET, C#.
        pattern = rf"(?<![A-Za-z0-9]){re.escape(skill)}(?![A-Za-z0-9])"
        if re.search(pattern, text, re.I):
            found.append(skill)
    return sorted(set(found))


def _find_name(text: str) -> str:
    """First plausible line. Deliberately low confidence.

    Resumes have no fixed layout: a name may be in a header, a text box, or
    absent from the text layer entirely. This is a suggestion for the
    candidate to correct, which is why it is never applied automatically.
    """
    for raw in text.splitlines()[:8]:
        line = raw.strip()
        if not (2 <= len(line.split()) <= 4):
            continue
        if NOT_A_NAME.search(line):
            continue
        if len(line) > 60:
            continue
        words = line.split()
        if all(w[:1].isupper() for w in words if w):
            return line
    return ""


def _years_of_experience(text: str) -> float:
    """Prefer a stated figure; otherwise add up date ranges.

    Summing ranges over-counts overlapping roles and under-counts gaps, which
    is exactly why this carries a low confidence and is shown for correction.
    """
    stated = STATED_YEARS_RE.search(text)
    if stated:
        value = int(stated.group(1))
        if 0 < value <= 50:
            return float(value)

    this_year = date.today().year
    total = 0
    for start, end in YEAR_RANGE_RE.findall(text):
        start_year = int(start)
        end_year = this_year if not end.isdigit() else int(end)
        if 0 < end_year - start_year <= 45:
            total += end_year - start_year
    return float(min(total, 45))


def _find_all(text: str, vocabulary: list[str]) -> list[str]:
    return sorted({item for item in vocabulary if re.search(re.escape(item), text, re.I)})


def parse_resume(data: bytes, content_type: str) -> dict[str, tuple[object, float]]:
    """Returns {field: (value, confidence)}.

    Only fields actually found are included — an absent field is absent, not
    an empty string presented as a finding.
    """
    text = extract_text(data, content_type)
    if not text.strip():
        # Almost always a scanned image with no text layer.
        raise ParseError(
            "No text could be read from this file. If it is a scan, upload a "
            "text-based PDF or DOCX instead."
        )

    results: dict[str, tuple[object, float]] = {}

    if email := EMAIL_RE.search(text):
        results["email"] = (email.group(0).lower(), 0.95)

    if phone := PHONE_RE.search(text):
        results["phone"] = (f"+1{phone.group(1)}{phone.group(2)}{phone.group(3)}", 0.95)

    if skills := _find_skills(text):
        results["skills"] = (skills, 0.80)

    if education := _find_all(text, DEGREES):
        results["education"] = (education, 0.70)

    if certifications := _find_all(text, CERTIFICATIONS):
        results["certifications"] = (certifications, 0.70)

    if years := _years_of_experience(text):
        results["years_experience"] = (years, 0.55)

    if name := _find_name(text):
        results["full_name"] = (name, 0.40)

    if titles := _find_all(text, TITLE_WORDS):
        results["job_titles"] = (titles, 0.35)

    return results


def experience_band(years: float) -> str:
    """Map to the Candidate.experience_level choices."""
    if years <= 2:
        return "0-2"
    if years <= 5:
        return "3-5"
    if years <= 10:
        return "6-10"
    return "10+"
